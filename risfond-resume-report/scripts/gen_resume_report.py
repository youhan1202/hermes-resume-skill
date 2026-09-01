"""锐仕方达简历报告生成脚本模板（2026-09 页眉页脚版）
对齐规范：①页眉左对齐logo图片 ②正文首行居中"简历报告"标题图
③页脚：公司名+地址+电话 ④基本信息4行×2列无边框表格
⑤中文宋体+西文Times New Roman ⑥项目经历板块
使用：复制本脚本到工作目录，修改候选人信息区后 python gen_resume_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

doc = Document()

BLUE = RGBColor(0x2E, 0x75, 0xB5)
BLACK = RGBColor(0, 0, 0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'   # 西文/数字
style.font.size = Pt(10.5)
style.font.color.rgb = BLACK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 中文
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.5

# ===== 页面设置：页边距 + 页眉页脚 =====
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    section.header_distance = Cm(1.0)   # 页眉距边界
    section.footer_distance = Cm(1.0)   # 页脚距边界

# ===== 页眉：左对齐 logo 图片 =====
ASSETS = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'assets')
HEADER_LOGO = os.path.join(ASSETS, 'header_logo.png')

section = doc.sections[0]
header = section.header
header.is_linked_to_previous = False

if os.path.exists(HEADER_LOGO):
    # 清空默认段落
    for p in header.paragraphs:
        p.clear()
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = hp.add_run()
    run.add_picture(HEADER_LOGO, width=Cm(6))   # 页眉logo宽度6cm，保持纵横比
else:
    # 无图片时用文字
    hp = header.paragraphs[0]
    hp.clear()
    R_h = hp.add_run('锐仕方达 RISFOND')
    R_h.bold = True
    R_h.font.size = Pt(14)
    R_h.font.color.rgb = RGBColor(0x00, 0x50, 0x8E)
    R_h.font.name = 'Times New Roman'
    R_h.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ===== 页脚：公司名 + 地址 + 电话 =====
footer = section.footer
footer.is_linked_to_previous = False

# 清空默认段落
for p in footer.paragraphs:
    p.clear()

# 页脚内容（三行居中）
footer_lines = [
    '公司：锐仕方达人才科技集团有限公司常州第一分公司',
    '地址：常州市武进区湖塘镇吾悦广场1幢12B03-12B05号',
    '公司电话：15295000586',
]

for i, line in enumerate(footer_lines):
    if i == 0:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fp.paragraph_format.space_before = Pt(0)
    fp.paragraph_format.space_after = Pt(0)
    fp.paragraph_format.line_spacing = 1.0
    run = fp.add_run(line)
    run.font.size = Pt(8)   # 页脚用小号字
    run.font.name = 'Times New Roman'
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)   # 灰色
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def R(p, text, bold=False, size=10.5, color=None):
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'   # 西文/数字
    run.font.color.rgb = color if color else BLACK
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')   # 中文
    return run


def S(text, bold=False, size=10.5, color=None, align=None):
    p = doc.add_paragraph()
    R(p, text, bold, size, color)
    if align is not None:
        p.alignment = align
    return p


def P(parts):
    """parts: [(text, bold, size, color), ...]"""
    p = doc.add_paragraph()
    for text, bold, size, color in parts:
        R(p, text, bold, size, color)
    return p


def NL():
    doc.add_paragraph()


# ===== 正文首行：顶部标题图居中 =====
TITLE_REPORT_IMG = os.path.join(ASSETS, 'title_report.png')

if os.path.exists(TITLE_REPORT_IMG):
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(12)
    run_title = p_title.add_run()
    run_title.add_picture(TITLE_REPORT_IMG, width=Cm(15))   # 标题图宽度15cm，保持纵横比
else:
    # 无图片时用文字
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    R(p_title, '简历报告', bold=True, size=18, color=RGBColor(0x00, 0x50, 0x8E))
    p_title.paragraph_format.space_after = Pt(12)


# ===== 模块标题图片 =====
TITLE_IMGS = {
    'basic': os.path.join(ASSETS, 'title_basic_info.png'),
    'edu': os.path.join(ASSETS, 'title_education.png'),
    'work': os.path.join(ASSETS, 'title_work_experience.png'),
    'project': os.path.join(ASSETS, 'title_project_experience.png'),
}


def add_title(key_or_path, fallback_text, width_cm=8.0):
    """添加模块标题：图片存在用图片，否则用【】文字（不加粗）"""
    path = TITLE_IMGS.get(key_or_path, key_or_path)
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.add_run().add_picture(path, width=Cm(width_cm))
    else:
        S(fallback_text)   # 不加粗


def add_basic_info_table(info):
    """基本信息 4行×2列 无边框表格。
    info 示例：[('姓名：张海默', '性别：男'), ('年龄：46岁', '学历：本科'),
               ('籍贯：江苏-苏州', '婚姻：已婚已育'), ('现居：上海', '')]"""
    tbl = doc.add_table(rows=len(info), cols=2)   # 默认样式无边框
    for i, (left, right) in enumerate(info):
        for j, text in enumerate((left, right)):
            cell = tbl.cell(i, j)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.5
            if text:
                R(p, text)
    return tbl


def add_company_intro(text):
    """公司介绍：标签加粗，内容不加粗"""
    p = doc.add_paragraph()
    R(p, '公司介绍：', bold=True)
    R(p, text)


def add_work(time_range, company, position):
    """工作经历标题行：蓝色加粗，三段\\t分隔"""
    P([(f'{time_range}\t{company}\t{position}', True, 10.5, BLUE)])


def add_project(time_range, proj_name, role, intro=None, duties=None, results=None):
    """项目经历段落：标题行蓝色加粗；项目介绍/职责/业绩标签加粗、内容不加粗"""
    P([(f'{time_range}\t{proj_name}\t{role}', True, 10.5, BLUE)])
    if intro:
        p = doc.add_paragraph()
        R(p, '项目介绍：', bold=True)
        R(p, intro)
    if duties:
        for k, d in enumerate(duties):
            p = doc.add_paragraph()
            if k == 0:
                R(p, '项目职责：', bold=True)
            R(p, d)
    if results:
        for k, d in enumerate(results):
            p = doc.add_paragraph()
            if k == 0:
                R(p, '项目业绩：', bold=True)
            R(p, d)
    NL()


# ===== 请修改以下候选人信息 =====
POSITION = "智算销售总监（国内）"   # 推荐职位（可省略）
CANDIDATE_NAME = "张海默"
BASIC_INFO = [
    ('姓名：张海默', '性别：男'),
    ('年龄：46岁', '学历：本科'),
    ('籍贯：江苏-苏州', '婚姻：已婚已育'),
    ('现居：上海', ''),
]
# 职业状态含到岗周期（基本信息表格无到岗周期行）
CAREER_STATUS = '离职，1个月内可尽快到岗。'

# 职业经历（时间倒序，时长紧跟时间）
EXPERIENCES = [
    "2021/07-2026/07 (5年1个月)  超云数字技术集团有限公司 大客户销售管理",
    "2019/09-2021/07 (1年11个月) 浪潮商用 互联网行业部副总经理",
]

# 核心优势（已确认版本；业绩丰富时第3条可用「相关业绩」子结构）
ADVANTAGES = [
    "1、人选本科毕业于北京化工大学（211、双一流）电子信息工程专业。",
    "2、人选拥有22年IT/智算领域B2B大客户销售经验……",
    "3、相关业绩：",
    "3.1 超云-任职近三年年均销售额7000万、单笔订单2000-3000万，促成阿里存储项目签约1亿+、整体框架5亿+；",
    "3.2 蜂盒科技-期间操盘阿里云GPU/存储项目年中标额超5亿；",
    "累计操盘项目金额超20亿，客户覆盖阿里、腾讯、科大讯飞等头部企业。",
    "4、人选具备从0到1开拓区域市场、搭建销售体系的能力……",
]

EDUCATION_ITEMS = [
    "1998/09-2002/07  北京化工大学（211）  电子信息工程  本科(统招)",
]

# ===== 生成文档（以下一般无需修改）=====

# 推荐信息
P([('推荐职位：' + POSITION, True, 10.5, None), ('\t推荐顾问：Alfred', True, 10.5, None)])
S('推荐日期：' + __import__('datetime').datetime.now().strftime('%Y-%m-%d'), bold=True)

# 基本信息：4行×2列无边框表格（标题图片可选）
add_title('basic', '')   # 基本信息标题图片（如样板）；不需要可注释本行
add_basic_info_table(BASIC_INFO)
NL()

# 综合评估 - 所有【】标签不加粗
S('综合评估：', bold=True)
S('【职业经历】')
for e in EXPERIENCES:
    S(e)
S('【核心优势】')
for a in ADVANTAGES:
    S(a)
S('【职业状态】' + CAREER_STATUS)
S('【家庭情况】已婚已育，定居苏州。')
S('【薪资结构】')
S('目前薪资：33k*12薪=39.6w左右（奖金期权未兑现）')
S('期望薪资：面议')
S('【意向度】高')
NL()

# 教育背景
add_title('edu', '【教育背景】')
for e in EDUCATION_ITEMS:
    S(e)
NL()

# 工作经历
add_title('work', '【工作经历】')
add_work('2021/07-2026/07(5年1个月)', '超云数字技术集团有限公司', '大客户销售管理')
add_company_intro('成立于2010年，隶属于中国电子信息产业集团……（50-100字，搜索补充）')
P([('汇报对象：', True, 10.5, BLACK), ('销售副总裁', False, 10.5, BLACK)])
P([('下属人数：', True, 10.5, BLACK), ('7人', False, 10.5, BLACK)])
P([('工作职责：', True, 10.5, None), ('1、负责……', False, 10.5, None)])
P([('工作业绩：', True, 10.5, None), ('1、年均销售额7000万……', False, 10.5, None)])
P([('离职原因：', True, 10.5, None), ('公司管理层变动，业务线发展空间受限。', False, 10.5, None)])
NL()

# 项目经历（销售/项目型岗位默认需要；技术岗无项目信息可整段省略）
add_title('project', '【项目经历】')
add_project(
    '2024/04-2026/07', '国内某知名智算服务商智算项目', '销售及项目管理',
    intro='国内某知名智算服务商智算中心建设项目，竞争智算项目建设硬件标的。',
    duties=['1、全面管理项目从立项到交付的全链条工作；', '2、击败竞争对手，高价中标，定义项目方向和后续目标。'],
    results=['1、公司首个智算中心推理集群建设，树立行业标杆；', '2、完成几千万级别订单，开拓新业务方向。'],
)

# 尾部
S('以上是我公司对候选人' + CANDIDATE_NAME + '的调研分析报告，文件中有需要保密的内容，请传阅时严格控制在与此次招聘有关的负责人手中。此文件的内容是我公司与候选人之间的面试总结。请贵公司根据人才情况和公司实际情况给予其相应待遇。')
S('调研单位：锐仕方达人才科技集团有限公司常州第一分公司')

output = f'锐仕方达-{CANDIDATE_NAME}-{POSITION}.docx'
doc.save(output)
print(f'DONE: {output}')
