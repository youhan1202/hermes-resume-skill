"""
锐仕方达简历报告生成脚本模板
使用：python gen_resume_report.py
生成前修改下方的候选人信息、工作经历等数据。
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

doc = Document()

BLUE = RGBColor(0x2E, 0x75, 0xB5)
BLACK = RGBColor(0, 0, 0)

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(10.5)
style.font.color.rgb = BLACK
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def R(p, text, bold=False, size=10.5, color=None):
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run.font.color.rgb = color if color else BLACK
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return run

def S(doc, text, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    R(p, text, bold, size, color)
    return p

def P(doc, parts):
    """parts: [(text, bold, size, color), ...]"""
    p = doc.add_paragraph()
    for text, bold, size, color in parts:
        R(p, text, bold, size, color)
    return p

def NL(doc):
    doc.add_paragraph()

def add_job(doc, time_span, company, title, intro, report_to, subordinates, duties, achievements=None, leave=None):
    """添加一段工作经历"""
    P(doc, [(f'{time_span}\t{company}\t{title}', True, 10.5, BLUE)])
    P(doc, [('公司介绍：', True, 10.5, None), (intro, False, 10.5, None)])
    P(doc, [('汇报对象：', True, 10.5, BLACK), (report_to, False, 10.5, None)])
    if subordinates:
        P(doc, [('下属人数：', True, 10.5, BLACK), (subordinates, False, 10.5, None)])
    S(doc, '工作职责：', bold=True)
    for d in duties:
        S(doc, d)
    if achievements:
        S(doc, '工作业绩：', bold=True)
        for a in achievements:
            S(doc, a)
    if leave:
        S(doc, leave)
    NL(doc)

# ===== 请修改以下候选人信息 =====
POSITION = "电气工程师"      # 推荐职位
CANDIDATE_NAME = "张三"      # 候选人姓名
GENDER = "男"
AGE = "30岁"
EDUCATION = "本科"
NATIVE = "江苏-南京"
MARRIAGE = "已婚已育"
LOCATION = "江苏-南京"
AVAILABILITY = "1个月"
SALARY_CURRENT = "15k*13薪=19.5w税前"
SALARY_EXPECTED = "18k左右"
INTENTION = "高"

# 职业经历（时间倒序）
EXPERIENCES = [
    "2022/01-至今 (4年6个月)  XX公司  电气工程师",
    "2019/03-2021/12 (2年9个月)  YY公司  电气工程师",
]

# 核心优势（统招985/211优先写学历，非统本/专科不写学历）
ADVANTAGES = [
    "1、人选本科毕业于XX大学XX专业，拥有X年电气工程领域实战经验。",
    "2、人选精通…",
    "3、…",
]

# 教育背景
EDUCATION_ITEMS = [
    " 2014/09-2018/06  大学名称  专业  本科(统招)",
]

# 工作经历详情
JOBS = [
    {
        "time": "2022/01-至今(4年6个月)", "company": "XX公司", "title": "电气工程师",
        "intro": "公司主营业务介绍。",
        "report_to": "电气主管", "subordinates": None,
        "duties": [
            "1、工作职责1",
            "2、工作职责2",
        ],
        "achievements": None,
        "leave": "离职原因：XXX",
    },
]

# ===== 生成文档（以下一般无需修改） =====
# 标题
P(doc, [('推荐职位：' + POSITION, True, 12, None), ('    ', False, 12, None), ('推荐顾问：Alfred', True, 12, None)])
S(doc, '推荐日期：' + __import__('datetime').datetime.now().strftime('%Y-%m-%d'), bold=True)
NL(doc)

# 基本信息
for line in ['姓名：' + CANDIDATE_NAME, '性别：' + GENDER, '年龄：' + AGE, '学历：' + EDUCATION,
             '籍贯：' + NATIVE, '婚姻：' + MARRIAGE, '现居：' + LOCATION, '到岗周期：' + AVAILABILITY]:
    S(doc, line)
NL(doc)

# 综合评估
S(doc, '综合评估：', bold=True)
NL(doc)
S(doc, '【职业经历】')
for e in EXPERIENCES:
    S(doc, e)
NL(doc)
S(doc, '【核心优势】')
for a in ADVANTAGES:
    S(doc, a)
NL(doc)

S(doc, '【职业状态】'); S(doc, f'{"已离职" if "离职" in str(AVAILABILITY) else "在职"}，到岗周期{AVAILABILITY}。')
S(doc, '【家庭情况】'); S(doc, f'{MARRIAGE}。')
S(doc, '【薪资结构】'); S(doc, f'目前薪资：{SALARY_CURRENT}'); S(doc, f'期望薪资：{SALARY_EXPECTED}')
S(doc, '【意向度】' + INTENTION)
S(doc, '【面试安排】提前沟通')
NL(doc)

# 教育
for e in EDUCATION_ITEMS:
    S(doc, e)
NL(doc)

# 工作经历
S(doc, '工作经历：', bold=True)
NL(doc)
for j in JOBS:
    add_job(doc, j["time"], j["company"], j["title"], j["intro"], j["report_to"], j["subordinates"], j["duties"], j.get("achievements"), j.get("leave"))

# 尾部
NL(doc)
S(doc, f'以上是我公司对候选人{CANDIDATE_NAME}的调研分析报告，文件中有需要保密的内容，请传阅时严格控制在与此次招聘有关的负责人手中。此文件的内容是我公司与候选人之间的面试总结。请贵公司根据人才情况和公司实际情况给予其相应待遇。')
NL(doc)
S(doc, '调研单位：锐仕方达人才科技集团有限公司常州第一分公司')

output = f'锐仕方达-{CANDIDATE_NAME}-{POSITION}.docx'
doc.save(output)
print(f'DONE: {output}')
